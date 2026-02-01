"""
Document Processing Service for Extracting Text from Various File Formats.

This module handles the extraction of text content from uploaded documents,
specifically PDF, DOCX, and plain text files. It's the first step in the
document processing pipeline - converting binary file formats into clean text
that can be chunked and embedded for semantic search.

The service supports:
- PDF files (using PyPDF2)
- DOCX/DOC files (using python-docx)
- Plain text files (UTF-8 or Latin-1 encoding)

All extracted text is cleaned and normalized to ensure consistent processing
in downstream steps.
"""

import logging
from dataclasses import dataclass
from typing import BinaryIO

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.utils.helpers import get_file_extension

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """
    Data class representing a processed document with extracted content.

    This structure holds the result of document processing, including:
    - The extracted text content
    - Original filename for reference
    - Document type (resume or job_posting) for categorization
    - Page count for user feedback
    - Word count calculated automatically after initialization
    """

    content: str
    filename: str
    doc_type: str  # "resume" or "job_posting"
    page_count: int = 1
    word_count: int = 0

    def __post_init__(self) -> None:
        """
        Calculate word count after object initialization.

        This runs automatically after the dataclass is created, ensuring
        word_count is always accurate based on the actual content.
        """
        self.word_count = len(self.content.split())


class DocumentProcessor:
    """
    Document Processor for Extracting Text from Various File Formats.

    This class provides methods to extract text content from PDF, DOCX, and
    plain text files. It handles different file formats by delegating to
    format-specific extraction methods, then cleans and normalizes the
    extracted text for consistent downstream processing.
    """

    # Supported file extensions that users can upload
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

    def process(
        self,
        file: BinaryIO,
        filename: str,
        doc_type: str,
    ) -> ProcessedDocument:
        """
        Process an uploaded file and extract its text content.

        This is the main entry point for document processing. It:
        1. Validates the file extension
        2. Routes to the appropriate extraction method based on file type
        3. Cleans and normalizes the extracted text
        4. Returns a ProcessedDocument with all metadata

        Args:
            file: Binary file object from Streamlit file uploader.
            filename: Original filename for reference and display.
            doc_type: Document type ("resume" or "job_posting") for categorization.

        Returns:
            ProcessedDocument containing extracted text and metadata.

        Raises:
            ValueError: If file type is not in the supported extensions list.
        """
        extension = get_file_extension(filename)

        # Validate file type before processing
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {extension}. "
                f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        logger.info(f"Processing {doc_type}: {filename}")

        # Route to appropriate extraction method based on file extension
        if extension == ".pdf":
            content, page_count = self._extract_pdf(file)
        elif extension in {".docx", ".doc"}:
            content, page_count = self._extract_docx(file)
        elif extension == ".txt":
            content, page_count = self._extract_txt(file)
        else:
            # This shouldn't happen due to validation above, but included for safety
            raise ValueError(f"Unsupported file type: {extension}")

        # Clean up extracted text to normalize whitespace and formatting
        # This ensures consistent processing regardless of source format
        content = self._clean_text(content)

        logger.info(
            f"Extracted {len(content)} characters from {filename} "
            f"({page_count} pages, {len(content.split())} words)"
        )

        return ProcessedDocument(
            content=content,
            filename=filename,
            doc_type=doc_type,
            page_count=page_count,
        )

    # =============================================================================
    # FORMAT-SPECIFIC EXTRACTION METHODS
    # =============================================================================

    def _extract_pdf(self, file: BinaryIO) -> tuple[str, int]:
        """
        Extract text content from a PDF file.

        PDFs are processed page-by-page using PyPDF2. We extract text from
        each page and join them with double newlines to preserve page boundaries.
        Empty pages are skipped to avoid adding blank content.

        Args:
            file: Binary file object containing PDF data.

        Returns:
            Tuple of (extracted text, actual page count from PDF).
        """
        reader = PdfReader(file)
        pages = []

        # Extract text from each page
        for page in reader.pages:
            text = page.extract_text()
            # Only add non-empty pages to avoid blank content
            if text:
                pages.append(text)

        # Join pages with double newline to preserve page separation
        return "\n\n".join(pages), len(reader.pages)

    def _extract_docx(self, file: BinaryIO) -> tuple[str, int]:
        """
        Extract text content from a DOCX or DOC file.

        DOCX files contain structured content (paragraphs and tables). We extract:
        1. All paragraph text
        2. Table content (converted to pipe-separated format for readability)

        Since DOCX doesn't have explicit page breaks like PDFs, we estimate
        page count based on character count (approximately 3000 chars per page).

        Args:
            file: Binary file object containing DOCX/DOC data.

        Returns:
            Tuple of (extracted text, estimated page count).
        """
        doc = DocxDocument(file)
        paragraphs = []

        # Extract text from all paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                paragraphs.append(para.text)

        # Also extract text from tables (resumes often use tables for formatting)
        # We convert table rows to pipe-separated text for readability
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        content = "\n\n".join(paragraphs)
        # Estimate page count: approximately 3000 characters per page
        # We use max(1, ...) to ensure at least 1 page even for very short documents
        estimated_pages = max(1, len(content) // 3000)

        return content, estimated_pages

    def _extract_txt(self, file: BinaryIO) -> tuple[str, int]:
        """
        Extract text content from a plain text file.

        Plain text files are the simplest - we just need to read and decode
        the binary content. We try UTF-8 first (most common), then fall back
        to Latin-1 if UTF-8 decoding fails.

        Args:
            file: Binary file object containing plain text data.

        Returns:
            Tuple of (text content, page count of 1 since text files don't have pages).
        """
        content = file.read()

        # Try UTF-8 first (most common encoding), fallback to Latin-1 if it fails
        # This handles edge cases where files might be in older encodings
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        return text, 1

    # =============================================================================
    # TEXT CLEANING AND NORMALIZATION
    # =============================================================================

    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text for consistent processing.

        Different file formats and extraction methods can produce inconsistent
        whitespace. This method normalizes:
        - Multiple spaces/tabs → single space
        - Multiple newlines (3+) → double newline (preserves paragraph breaks)
        - Leading/trailing whitespace → removed

        This ensures downstream processing (chunking, embedding) works
        consistently regardless of source format.

        Args:
            text: Raw extracted text with potentially inconsistent whitespace.

        Returns:
            Cleaned text with normalized whitespace.
        """
        import re

        # Replace multiple spaces or tabs with a single space
        text = re.sub(r"[ \t]+", " ", text)
        # Replace 3+ consecutive newlines with double newline
        # This preserves paragraph breaks while removing excessive blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Remove leading and trailing whitespace
        return text.strip()

    # =============================================================================
    # TEXT INPUT PROCESSING (FOR PASTED CONTENT)
    # =============================================================================

    def process_text(self, text: str, title: str, doc_type: str) -> ProcessedDocument:
        """
        Process raw text input directly (for job postings pasted as text).

        This method handles the case where users paste job posting text directly
        into the UI instead of uploading a file. It applies the same cleaning
        logic as file-based processing to ensure consistency.

        Args:
            text: Raw text content pasted by the user.
            title: Title/name for the document (e.g., "Senior ML Engineer @ Google").
            doc_type: Document type ("resume" or "job_posting").

        Returns:
            ProcessedDocument with cleaned text content and metadata.
        """
        # Apply same cleaning logic as file-based processing
        content = self._clean_text(text)
        return ProcessedDocument(
            content=content,
            filename=title,
            doc_type=doc_type,
            page_count=1,  # Pasted text has no page concept
        )
